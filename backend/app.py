import os
import uuid
import json
import asyncio
from typing import Dict
from dotenv import load_dotenv
load_dotenv()

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
from groq import Groq

from backend.models.schemas import ResearchRequest, ResearchResponse
from backend.graph.build_graph import research_app
from backend.tools.vector_store import vector_store
from pypdf import PdfReader

app = FastAPI(title="ResearchMind AI Gateway", version="1.0.0")


app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# In-memory session reports store
session_reports: Dict[str, dict] = {}
active_connections: Dict[str, WebSocket] = {}

@app.get("/")
def read_root():
    return {"status": "online", "app": "ResearchMind AI API Gateway"}

@app.post("/research", response_model=ResearchResponse)
def start_research(req: ResearchRequest):
    session_id = str(uuid.uuid4())
    session_reports[session_id] = {
        "query": req.query,
        "status": "initialized",
        "final_report": None,
        "citations": []
    }
    return ResearchResponse(session_id=session_id, message="Research session created successfully.")

@app.websocket("/ws/{session_id}")
async def websocket_research(websocket: WebSocket, session_id: str):
    await websocket.accept()
    active_connections[session_id] = websocket
    
    session_data = session_reports.get(session_id, {"query": "Solid-State Battery Breakthroughs 2026"})
    query = session_data.get("query", "Solid-State Battery Breakthroughs 2026")

    initial_state = {
        "query": query,
        "session_id": session_id,
        "subtasks": [],
        "research_notes": [],
        "verified_notes": [],
        "outline": [],
        "draft_report": "",
        "final_report": "",
        "citations": [],
        "confidence_score": 1.0,
        "retry_count": 0,
        "status_log": []
    }

    try:
        # Stream graph state execution
        for output in research_app.stream(initial_state):
            for node_name, state_update in output.items():
                status_log = state_update.get("status_log", [])
                latest_event = status_log[-1] if status_log else {"agent": node_name, "status": "completed", "message": f"{node_name} finished."}
                
                await websocket.send_text(json.dumps({
                    "type": "agent_event",
                    "agent": node_name,
                    "event": latest_event,
                    "confidence_score": state_update.get("confidence_score", 1.0)
                }))
                await asyncio.sleep(0.3)
                
                if "final_report" in state_update and state_update["final_report"]:
                    session_reports[session_id] = {
                        "query": query,
                        "status": "completed",
                        "final_report": state_update["final_report"],
                        "citations": state_update.get("citations", [])
                    }
                    await websocket.send_text(json.dumps({
                        "type": "done",
                        "final_report": state_update["final_report"],
                        "citations": state_update.get("citations", [])
                    }))

    except WebSocketDisconnect:
        print(f"Client disconnected for session {session_id}")
    except Exception as e:
        print(f"WebSocket error: {e}")
        await websocket.send_text(json.dumps({"type": "error", "message": str(e)}))
    finally:
        if session_id in active_connections:
            del active_connections[session_id]

@app.post("/documents/upload")
async def upload_pdf(session_id: str = Form(...), file: UploadFile = File(...)):
    if not file.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Only PDF files are supported.")
    
    try:
        reader = PdfReader(file.file)
        text = ""
        for page in reader.pages:
            t = page.extract_text()
            if t:
                text += t + "\n"

        # Simple chunking
        chunks = [text[i:i+800] for i in range(0, len(text), 700)]
        metadatas = [{"filename": file.filename, "session_id": session_id} for _ in chunks]

        vector_store.add_documents(session_id=session_id, chunks=chunks, metadatas=metadatas)
        return {"filename": file.filename, "chunks_indexed": len(chunks), "message": "PDF uploaded & indexed in ChromaDB."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/transcribe")
async def transcribe_audio(file: UploadFile = File(...)):
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY is not configured on backend.")
    
    try:
        content = await file.read()
        client = Groq(api_key=api_key)
        transcription = client.audio.transcriptions.create(
            file=(file.filename or "speech.webm", content),
            model="whisper-large-v3-turbo",
            response_format="json",
            language="en"
        )
        return {"text": transcription.text}
    except Exception as e:
        print(f"Groq transcription error: {e}")
        raise HTTPException(status_code=500, detail=f"Speech transcription failed: {str(e)}")

@app.get("/export/{session_id}")
def export_report(session_id: str, format: str = "md"):
    data = session_reports.get(session_id)
    if not data or not data.get("final_report"):
        raise HTTPException(status_code=404, detail="Report not found or research in progress.")
    
    report_text = data["final_report"]
    query = data.get("query", "Research Report")

    if format == "md":
        return Response(
            content=report_text,
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename=report_{session_id[:8]}.md"}
        )

    elif format == "pdf":
        try:
            import io
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.lib import colors
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
            from reportlab.lib.enums import TA_LEFT, TA_CENTER

            buf = io.BytesIO()
            doc = SimpleDocTemplate(
                buf, pagesize=A4,
                leftMargin=2.5*cm, rightMargin=2.5*cm,
                topMargin=2.5*cm, bottomMargin=2.5*cm
            )
            styles = getSampleStyleSheet()

            accent = colors.HexColor("#34D399")
            dark = colors.HexColor("#0B0D10")

            title_style = ParagraphStyle("Title", parent=styles["Heading1"], fontSize=22, textColor=accent, spaceAfter=8, leading=28)
            h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=15, textColor=colors.HexColor("#E5E7EB"), spaceAfter=6, spaceBefore=14)
            h3_style = ParagraphStyle("H3", parent=styles["Heading3"], fontSize=12, textColor=colors.HexColor("#9CA3AF"), spaceAfter=4, spaceBefore=10)
            body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10, textColor=colors.HexColor("#D1D5DB"), leading=16, spaceAfter=6)
            meta_style = ParagraphStyle("Meta", parent=styles["Normal"], fontSize=9, textColor=colors.HexColor("#6B7280"), leading=13)

            story = []
            story.append(Paragraph("ResearchMind AI", ParagraphStyle("Brand", parent=styles["Normal"], fontSize=10, textColor=accent)))
            story.append(Spacer(1, 6))
            story.append(Paragraph(f"Research Report: {query}", title_style))
            story.append(HRFlowable(width="100%", thickness=1, color=accent, spaceAfter=16))

            import re
            for line in report_text.split("\n"):
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 5))
                elif line.startswith("### "):
                    story.append(Paragraph(line[4:], h3_style))
                elif line.startswith("## "):
                    story.append(Paragraph(line[3:], h2_style))
                elif line.startswith("# "):
                    story.append(Paragraph(line[2:], title_style))
                elif line.startswith("- ") or line.startswith("* "):
                    story.append(Paragraph(f"• {line[2:]}", body_style))
                elif line.startswith("---"):
                    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#374151"), spaceAfter=8))
                else:
                    # Strip basic markdown bold/italic for PDF
                    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                    clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                    clean = re.sub(r'\[.*?\]\(.*?\)', '', clean)
                    story.append(Paragraph(clean, body_style))

            story.append(Spacer(1, 20))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#374151")))
            story.append(Spacer(1, 6))
            story.append(Paragraph("Generated by ResearchMind AI · Groq Llama 3.3 70B · Tavily/DuckDuckGo Search", meta_style))

            doc.build(story)
            buf.seek(0)
            return Response(
                content=buf.read(),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename=report_{session_id[:8]}.pdf"}
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"PDF generation failed: {str(e)}")

    elif format == "docx":
        try:
            import io
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            
            # Page margins
            for section in doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.2)
                section.right_margin = Inches(1.2)

            # Header branding
            hdr = doc.add_paragraph()
            hdr_run = hdr.add_run("ResearchMind AI  ·  AI Research Report")
            hdr_run.font.size = Pt(9)
            hdr_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
            hdr.paragraph_format.space_after = Pt(4)

            # Title
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(query)
            title_run.font.size = Pt(22)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0x34, 0xD3, 0x99)
            title_para.paragraph_format.space_after = Pt(12)

            doc.add_paragraph()

            import re
            for line in report_text.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                if line.startswith("### "):
                    p = doc.add_paragraph()
                    run = p.add_run(line[4:])
                    run.font.size = Pt(12)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
                elif line.startswith("## "):
                    p = doc.add_paragraph()
                    run = p.add_run(line[3:])
                    run.font.size = Pt(15)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0xE5, 0xE7, 0xEB)
                elif line.startswith("# "):
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:])
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x34, 0xD3, 0x99)
                elif line.startswith("- ") or line.startswith("* "):
                    p = doc.add_paragraph(style="List Bullet")
                    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])
                    clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                    clean = re.sub(r'\[.*?\]\(.*?\)', '', clean)
                    p.add_run(clean).font.size = Pt(10)
                else:
                    clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                    clean = re.sub(r'\*(.*?)\*', r'\1', clean)
                    clean = re.sub(r'\[.*?\]\(.*?\)', '', clean)
                    p = doc.add_paragraph()
                    p.add_run(clean).font.size = Pt(10)

            # Footer
            doc.add_paragraph()
            footer_para = doc.add_paragraph()
            footer_run = footer_para.add_run("Generated by ResearchMind AI · Groq Llama 3.3 70B · Tavily/DuckDuckGo Search")
            footer_run.font.size = Pt(8)
            footer_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return Response(
                content=buf.read(),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename=report_{session_id[:8]}.docx"}
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX generation failed: {str(e)}")

    else:
        return Response(content=report_text, media_type="text/plain")

